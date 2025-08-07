#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Oct 22 20:59:46 2019

@author: karthick
"""
import re
import sys
from fileinput import lineno

###Import all necessary packages
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import *
from docx.text.paragraph import Paragraph
from docx.text.paragraph import Run
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docxcompose.composer import Composer
from docx import Document as Document_compose
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
import io
import csv
import base64
from docx import Document
from docx.shared import Inches
import subprocess
import os
from docx.enum.text import WD_COLOR_INDEX

# Load the docx file into document object. You can input your own docx file in this step by changing the input path below:
document = Document(sys.argv[1])


# document = Document("Template1.docx")


##This function extracts the tables and paragraphs from the document object
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, doctwo):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# This function extracts the table from the document object as a dataframe
def read_docx_tables(tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    :
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """

    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    #    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in document.tables]
    else:
        try:
            return read_docx_tab(document.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise


# The combined_df dataframe will store all the content in document order including images, tables and paragraphs.
# If the content is an image or a table, it has to be referenced from image_df for images and table_list for tables using the corresponding image or table id that is stored in combined_df
# And if the content is paragraph, the paragraph text will be stored in combined_df
combined_df = pd.DataFrame(columns=['para_text', 'table_id', 'style'])
table_mod = pd.DataFrame(columns=['string_value', 'table_id'])

# The image_df will consist of base64 encoded image data of all the images in the document
image_df = pd.DataFrame(columns=['image_index', 'image_rID', 'image_filename', 'image_base64_string'])

# The table_list is a list consisting of all the tables in the document
table_list = []
xml_list = []

i = 0
imagecounter = 0

blockxmlstring = ''

for block in iter_block_items(document):
    print(str(block))
for block in iter_block_items(document):
    if 'text' in str(block):
        isappend = False

        runboldtext = ''
        for run in block.runs:
            if run.bold:
                runboldtext = runboldtext + run.text

        style = str(block.style.name)

        appendtxt = str(block.text)
        appendtxt = appendtxt.replace("\n", "")
        appendtxt = appendtxt.replace("\r", "")
        tabid = 'Novalue'
        paragraph_split = appendtxt.lower().split()

        isappend = True
        for run in block.runs:
            xmlstr = str(run.element.xml)
            my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
            print(my_namespaces)
            root = ET.fromstring(xmlstr)
            # print(my_namespaces)
            # Check if pic is there in the xml of the element. If yes, then extract the image data
            if 'pic:pic' in xmlstr:
                xml_list.append(xmlstr)
                for pic in root.findall('.//pic:pic', my_namespaces):
                    cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                    name_attr = cNvPr_elem.get("name")
                    blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                    embed_attr = blip_elem.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    isappend = True
                    appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                    document_part = document.part
                    image_part = document_part.related_parts[embed_attr]
                    image_base64 = base64.b64encode(image_part._blob)
                    image_base64 = image_base64.decode()
                    dftemp = pd.DataFrame(
                        {'image_index': [imagecounter], 'image_rID': [embed_attr], 'image_filename': [name_attr],
                         'image_base64_string': [image_base64]})
                    image_df = image_df._append(dftemp, sort=False)
                    style = 'Novalue'
                imagecounter = imagecounter + 1
    # print(str(block))
    elif 'table' in str(block):
        isappend = True
        style = 'Novalue'
        appendtxt = str(block)
        tabid = i
        dfs = read_docx_tables(tab_id=i)
        dftemp = pd.DataFrame({'para_text': [appendtxt], 'table_id': [i], 'style': [style]})
        table_mod = table_mod._append(dftemp, sort=False)
        table_list.append(dfs)
        i = i + 1
    if isappend:
        dftemp = pd.DataFrame({'para_text': [appendtxt], 'table_id': [tabid], 'style': [style]})
        combined_df = combined_df._append(dftemp, sort=False)


def hasImage(par):
    # """get all of the images in a paragraph
    # :param par: a paragraph object from docx
    # :return: a list of r:embed
    # """
    # ids = []
    # root = ET.fromstring(par._p.xml)
    # # namespace = {
    # #          'a':"http://schemas.openxmlformats.org/drawingml/2006/main", \
    # #          'r':"http://schemas.openxmlformats.org/officeDocument/2006/relationships", \
    # #          'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
    # namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas', 'cx': 'http://schemas.microsoft.com/office/drawing/2014/chartex', 'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006', 'o': 'urn:schemas-microsoft-com:office:office', 'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships', 'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math', 'v': 'urn:schemas-microsoft-com:vml', 'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing', 'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing', 'w10': 'urn:schemas-microsoft-com:office:word', 'w14': 'http://schemas.microsoft.com/office/word/2010/wordml', 'w15': 'http://schemas.microsoft.com/office/word/2012/wordml', 'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex', 'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup', 'wpi': 'http://schemas.microsoft.com/office/word/2010/wordprocessingInk', 'wne': 'http://schemas.microsoft.com/office/word/2006/wordml', 'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape', 'a': 'http://schemas.openxmlformats.org/drawingml/2006/main', 'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}
    # inlines = root.findall('.//wp:inline',namespace)
    # for inline in inlines:
    #     imgs = inline.findall('.//a:blip', namespace)
    #     for img in imgs:
    #         id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
    #     ids.append(id)
    #
    # return ids
    # imagecounter = 0
    ids = []
    for run in par.runs:
        xmlstr = str(run.element.xml)
        my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
        print(my_namespaces)
        root = ET.fromstring(xmlstr)
        # print(my_namespaces)
        # Check if pic is there in the xml of the element. If yes, then extract the image data
        if 'pic:pic' in xmlstr:
            for pic in root.findall('.//pic:pic', my_namespaces):
                cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                name_attr = cNvPr_elem.get("name")
                blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                embed_attr = blip_elem.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                ids.append(embed_attr)
    return ids


def find_header(header_file, text):
    with open(header_file, encoding='utf-8') as header:
        header = [x.strip().split("\n")[0] for x in header if x.strip() != '']
    for line in header:
        if re.search(line,text.lower()):
            return line

    # print(list)
    # if list.__len__() > 0:
    #     return True


def return_image(rID):
    # if have image_df.iloc[0] , iloc[1] then can guess the iamge_df.iloc will return an array.
    for iloc in image_df.iloc:
        if iloc["image_rID"] == str(rID):
            image_base64_fromDF = iloc["image_base64_string"]
            return image_base64_fromDF


combined_df = combined_df.reset_index(drop=True)
image_df = image_df.reset_index(drop=True)
# demo1
# DataFrames array, iloc extract, code and decode the base64. Create a .docx file and add a picture into. The attribute of picture/
# merge 2 list still remain order : not finished

# print("image in between the text")
# print(combined_df)
# new_document =  Document()
# print(combined_df.iloc[0]["para_text"])
# print(combined_df.iloc[1]["para_text"])
# # print(combined_df.iloc[0]["image_base64_string"])
# # print(combined_df.iloc[0]["image_base64_string"].decode())
# # new_document.add_picture()
# print("---")
# print(image_df.iloc[0]["image_filename"])
# # print(image_df.iloc[0]["image_base64_string"].decode())
# print(base64.b64decode(image_df.iloc[0]["image_base64_string"]))

# # my_pic2 = bytes(my_pic)
# # print(type(my_pic2))
# new_document.add_paragraph("hehe")
# print(type(my_pic))

# new_document.save("hehe.docx")
#
# # use kdiff to merge the combine_df and the text frames from paragraphs.
# list1 = ["Document1/rID1","haha","Document1/rID2","hehe","huhu"]
# list2 = ["olala","haha","olele","hehe","huhu"]
# combined_set = set(list1).union(list2)
# print(combined_set)
# # {'olala', 'hehe', 'huhu', 'olele', 'Document1/rID2', 'haha', 'Document1/rID1'}

# expected: Document1/rID1 , olala , haha , Document1/rID2, olele, hehe, huhu

# demo2 copy the same file content from the original .docx with the same order.
# newdocument=Document()
# for para in document.paragraphs:
#     newdocument.add_paragraph(para.text)
#     if hasImage(para) != "None":
#         for rID in hasImage(para):
#             my_pic = base64.b64decode(return_image(rID))
#             image_stream = io.BytesIO(initial_bytes=my_pic)
#             newdocument.add_picture(image_stream,width=Inches(3))
# newdocument.save("hehe.docx")

# demo3
f = open("input.txt", "w", encoding='utf-8')
for para in document.paragraphs:
    print(type(para.text))
    # this is the same with .strip()
    # f.writelines(str(para.text).strip('\n') + "\n")

    # only write down the line with no empty, else we dont action (mean dont write anything)
    if re.match(r'^(?!\s*$)', para.text):
        f.writelines(str(para.text).strip() + "\n")
    for rID in hasImage(para):
        f.writelines(rID + "\n")
f.close()

# demo4
# use vim command or subprocess to remove dupplicate empty lines.
# use the .exe command from window to run for arrange
# read from the new text, if the text line is normal then no need to edit, but it has the regex form ^rID(1 or more number) then extract that line and input into the subfunct return image has created.
# have order
# subprocess.run("demo_console.exe > output.txt", shell="true")

# put the output arrange of demo_console in later of note to not effect the page number
# os.system("demo_console.exe header.txt input.txt output.txt > output/%s.txt"%(str((sys.argv[1].strip('.docx').strip('input\/')))))
os.system("demo_console.exe header.txt input.txt output.txt")

# no dupplicate line
subprocess.run("uniq.exe output.txt output2.txt", shell="true")
# no empty line
# subprocess.run("sed -i '/^$/d' output2.txt", shell = "true")   why this command erase all lines of .txt ?
# os.system("sed -i '/^$/d' output2.txt") this command also do that
# os.system("sed -i '/^$/d' output2.txt")# this command also do that
# the word file and text file have the same this format
# subprocess.run("cp -rf output2.txt output.txt", shell = "true")


# subprocess.run("vim.exe -c '%!uniq' output.txt", shell = "true") when run this command, it open the output.txt after finish but the file is told that readonly and need to add :w! to save the new change.
# os.system("hehe.exe output.txt output2.txt") this command also success.
# because the road up and down are the same, so to understand what kind of data we should write into a file, we can get the type of stream of read.

# demo5
# by using this way, the text will loss the color gray from review, so how I transform the text and these image with some tags to text files, this will make them have the attribute and after we extract out again, we can have the color or know this is an image.
# use output2 to gen the .docx with picture.
# there is no time to review note on Pc so will remove the no_rid.docx
# final_pc=Document()
final_phone = Document()
f = open("output2.txt", "r", encoding='utf-8')
for line in f.readlines():
    if re.search('^rId', line):
        my_pic = base64.b64decode(return_image(line.strip()))
        image_stream = io.BytesIO(initial_bytes=my_pic)
        # final_pc.add_picture(image_stream,width=Inches(6.25))
        final_phone.add_picture(image_stream, width=Inches(6.25))
        # dont want to have empty line after rId
        final_phone.add_paragraph(line.strip())
        print("image", line)
    elif find_header("header.txt", line.strip()) != None:
        # final_pc.add_paragraph(line.strip())
        # final_phone.add_paragraph(line.strip())
        print("true")
        final_phone.add_paragraph(
        ).add_run(
            line.strip()
        ).font.highlight_color = WD_COLOR_INDEX.YELLOW
        print(line)
    else:
        final_phone.add_paragraph(line.strip())
        # final_phone.add_paragraph().add_run(line.strip()).font.highlight_color = WD_COLOR_INDEX.BLACK
        print(line)
f.close()
# print ("output/" + (sys.argv[1].strip('.docx')) + "_no_rid.docx")
# final_pc.save("output/" + (sys.argv[1].strip('.docx').strip('input\/')) + "_no_rid.docx")
# subprocess.run("cp -f ./template.docx ")
final_phone.save("output/" + (sys.argv[1].strip('.docx').strip('input\/')) + "_rid.docx")

os.system("cat output2.txt > output/%s.txt" % (str((sys.argv[1].strip('.docx').strip('input\/')))))
os.system("demo_console.exe header.txt input.txt output.txt >> output/%s.txt" % (
    str((sys.argv[1].strip('.docx').strip('input\/')))))
subprocess.run("rm -f output.txt", shell="true")
subprocess.run("rm -f output2.txt", shell="true")
subprocess.run("rm -f input.txt", shell="true")
